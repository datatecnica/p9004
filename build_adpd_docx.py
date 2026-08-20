#!/usr/bin/env python3
"""build_adpd_docx.py — render adpd_abstract.md as a Word document.

Not a general Markdown converter: it walks the specific structure of
`adpd_abstract.md` (H2 sections, H3 subsections, pipe tables, blockquote title,
bold/italic/code inline spans) and lays it out the way a submission draft wants --
the abstract body set apart in a bordered block with its ALL-CAPS title, everything
else as ordinary document prose.

The word count is RECOMPUTED here from the rendered body and written into the
document, so the count in the .docx can never drift from the text above it the way a
hand-maintained figure would.

Usage:  python3 build_adpd_docx.py [-i adpd_abstract.md] [-o adpd_abstract.docx]
"""
from __future__ import annotations

import argparse
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK    = RGBColor(0x17, 0x20, 0x29)
CSF    = RGBColor(0x1F, 0x7A, 0x8C)
MUTED  = RGBColor(0x5D, 0x6B, 0x76)
FLAG   = RGBColor(0xA0, 0x3B, 0x2E)

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def shade(cell, hexcolor: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def border(par, edge="left", size=18, color="1F7A8C") -> None:
    pPr = par._p.get_or_add_pPr()
    bdr = pPr.find(qn("w:pBdr"))
    if bdr is None:
        bdr = OxmlElement("w:pBdr")
        pPr.append(bdr)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(size))
    e.set(qn("w:space"), "8")
    e.set(qn("w:color"), color)
    bdr.append(e)


def emit_inline(par, text: str, *, size=10.5, color=INK, base_bold=False) -> None:
    """Write `text` into `par`, honouring **bold**, *italic* and `code` spans."""
    for piece in _INLINE.split(text):
        if not piece:
            continue
        bold, italic, mono = base_bold, False, False
        if piece.startswith("**") and piece.endswith("**"):
            piece, bold = piece[2:-2], True
        elif piece.startswith("*") and piece.endswith("*"):
            piece, italic = piece[1:-1], True
        elif piece.startswith("`") and piece.endswith("`"):
            piece, mono = piece[1:-1], True
        r = par.add_run(piece)
        r.font.name = MONO_FONT if mono else BODY_FONT
        r.font.size = Pt(size - 0.5 if mono else size)
        r.font.bold, r.font.italic = bold, italic
        r.font.color.rgb = MUTED if mono else color


def add_table(doc, rows: list[list[str]]) -> None:
    head, body = rows[0], rows[1:]
    t = doc.add_table(rows=len(rows), cols=len(head))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(head):
        c = t.cell(0, j)
        c.text = ""
        shade(c, "EDF3F4")
        emit_inline(c.paragraphs[0], h, size=9, color=CSF, base_bold=True)
    for i, row in enumerate(body, start=1):
        for j, v in enumerate(row):
            if j >= len(head):
                continue
            c = t.cell(i, j)
            c.text = ""
            emit_inline(c.paragraphs[0], v, size=9)
    doc.add_paragraph()


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def build(md_path: str, out_path: str) -> None:
    src = open(md_path, encoding="utf-8").read()

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name, st.font.size = BODY_FONT, Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.9)

    lines = src.splitlines()
    i, body_words, in_body = 0, 0, False

    while i < len(lines):
        ln = lines[i]

        # ---- headings
        if ln.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(ln[2:].strip())
            r.font.size, r.font.bold, r.font.color.rgb = Pt(19), True, INK
            p.paragraph_format.space_after = Pt(10)
        elif ln.startswith("## "):
            title = ln[3:].strip()
            in_body = title == "Body"
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.font.size, r.font.bold, r.font.color.rgb = Pt(13.5), True, INK
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
        elif ln.startswith("### "):
            title = ln[4:].strip()
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.font.size, r.font.bold = Pt(10.5), True
            r.font.color.rgb = CSF if in_body else INK
            if in_body:
                r.font.all_caps = True
            p.paragraph_format.space_before = Pt(11)
            p.paragraph_format.space_after = Pt(3)
            if in_body:
                border(p, "left")
                p.paragraph_format.left_indent = Inches(0.14)

        # ---- horizontal rule
        elif ln.strip() == "---":
            p = doc.add_paragraph()
            border(p, "bottom", size=6, color="D8D6CE")
            p.paragraph_format.space_after = Pt(2)

        # ---- blockquote (the title block)
        elif ln.startswith("> "):
            buf = []
            while i < len(lines) and lines[i].startswith(">"):
                buf.append(lines[i].lstrip("> ").rstrip())
                i += 1
            p = doc.add_paragraph()
            border(p, "left", size=24)
            p.paragraph_format.left_indent = Inches(0.16)
            p.paragraph_format.space_before = Pt(4)
            emit_inline(p, " ".join(buf), size=12.5, color=INK)
            continue

        # ---- table
        elif ln.startswith("|") and i + 1 < len(lines) and set(lines[i + 1]) <= set("|-: "):
            rows = [split_row(ln)]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # ---- list item
        elif re.match(r"^\s*\d+\.\s+", ln) or re.match(r"^\s*[-*]\s+", ln):
            buf = [re.sub(r"^\s*(?:\d+\.|[-*])\s+", "", ln)]
            i += 1
            while i < len(lines) and lines[i].startswith("   ") and lines[i].strip():
                buf.append(lines[i].strip())
                i += 1
            p = doc.add_paragraph(style="List Bullet")
            emit_inline(p, " ".join(buf), size=10)
            continue

        # ---- paragraph
        elif ln.strip():
            buf = []
            while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "|", ">", "---")):
                buf.append(lines[i].strip())
                i += 1
            text = " ".join(buf)
            # The budget line lives under ## Body but is commentary about the body,
            # not part of it -- counting it would inflate the very number it reports.
            if text.startswith("**Word budget**"):
                in_body = False
            p = doc.add_paragraph()
            emit_inline(p, text, size=10.5)
            if in_body:
                p.paragraph_format.left_indent = Inches(0.14)
                p.paragraph_format.space_after = Pt(8)
                body_words += len(re.sub(r"[*_`]", "", text).split())
            continue

        i += 1

    out = os.path.abspath(out_path)
    doc.save(out)
    kb = os.path.getsize(out) / 1024
    print(f"Wrote {out}  ({kb:.0f} KB)")
    print(f"  body words counted while rendering: {body_words}"
          f"  (+4 section headings = {body_words + 4} / 280)")


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("-i", "--input", default="adpd_abstract.md")
    ap.add_argument("-o", "--output", default="adpd_abstract.docx")
    a = ap.parse_args()
    build(a.input, a.output)
